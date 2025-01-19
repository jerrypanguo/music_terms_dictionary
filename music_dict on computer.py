import difflib
import pypinyin
import json
import os
from pypinyin import lazy_pinyin, Style
import random
import sys
import tty
import termios
import select
import urllib.request
import ssl
import time
from collections import Counter
import shutil
import zipfile

# 音乐术语字典
music_terms = {
    'A cappella': '无伴奏的声乐演唱',
    'A punta d\'arco': '用弓尖演奏',
    'Abbandonamente': '放任地，无拘束地',
    'Abbellimenti': '装饰音群',
    'Accarezzevole': '爱抚的，抚摸着',
    'Acciaccato': '重音的，突强的',
    'Accordatura': '调音，定音法',
    'Acustica': '声学的',
    'Affetto': '柔情的',
    'Agitato': '激动地，不宁地，惊慌地',
    'Agogica': '速度表情法',
    'Aggradevole': '妩媚的，令人怜爱的',
    'Al ponticello': '在琴桥附近演奏',
    'Al taco': '用弓根部演奏',
    'Alberti bass': '阿尔贝蒂低音，分解和弦伴奏音型',
    'Alla cadenza': '如华彩段般',
    'Alla corda': '靠近琴弦演奏',
    'Alla marcia': '进行曲风格',
    'Alla polacca': '波兰舞曲风格的',
    'Alla siciliana': '西西里舞曲风格的',
    'Alla zingara': '吉普赛风格的',
    'Angoscioso': '焦虑不安地，痛苦地',
    'Animato': '有生气的，活跃的',
    'Appoggiando': '倚音的',
    'Appoggiatura doppia': '双倚音',
    'Armonici artificiali': '人工泛音',
    'Arpeggiando': '琶音式的',
    'Arpeggiato': '分散和弦的',
    'Articolazione': '发音法，衔接法',
    
    'Bariolage': '快速交替演奏空弦和按弦',
    'Basso continuo': '通奏低音',
    'Basso ostinato': '固定低音，反复低音音型',
    'Battement': '装饰音，颤音',
    'Bicinium': '双声部曲，二重唱',
    'Bicordo': '双音',
    'Bisbigliando': '低语般的，极轻的颤音',
    'Bitonalità': '双调性',
    'Bizzarro': '古怪的',
    'Brillante': '华丽而灿烂的',
    'Buffo': '滑稽的',
    
    'Cadenza composta': '复合终止式',
    'Cadenza doppia': '双重终止式',
    'Cadenza evitata': '回避终止式',
    'Cadenza fiorita': '华彩终止式',
    'Cadenza plagale': '变格终止式',
    'Camminando': '流畅地，从容不迫地',
    'Cantabile': '如歌的',
    'Cantabile parlante': '如说话般的歌唱',
    'Canto fiorito': '华丽的歌唱风格',
    'Canto spianato': '平稳的歌唱风格',
    'Chiavette': '高位谱号系统',
    'Cluster': '音团，音簇',
    'Col legno battuto': '用弓杆敲击琴弦',
    'Col legno tratto': '用弓杆摩擦琴弦',
    'Coloratura': '花腔，装饰性的旋律段落',
    'Colpo d\'arco': '弓击音',
    'Combinazione di registri': '音栓组合',
    'Combinazione toni': '音的组合',
    'Con brio': '有精神，有活力的',
    'Con grazia': '优美的',
    'Contrappunto doppio': '双重对位法',
    'Contrappunto invertibile': '可转位对位法',
    'Contrappunto osservato': '严格对位法',
    'Corda vuota': '空弦',
    'A tempo': '速度还原',
    'accel.': '加速',
    'Accelerando': '加速的',
    'Adagio': '柔板；从容的；悠闲的',
    'Affettuoso': '深情的',
    'Al fine': '直到结束',
    'Allegretto': '小快板；稍快；比allegro稍慢的速度',
    'Allegro': '快板；欢快地；较活泼的速度',
    'Allegro molto': '非常快速的',
    'Andante': '行板；徐缓；行进的',
    'Andante con moto': '行进速度，带有动作感的',
    'Appassionato': '充满激情的',
    'Arioso': '空灵的，歌唱的',
    'Arpeggio': '琶音',
    'Attacca': '紧接着演奏',
    'Ben marcato': '清晰地强调',
    'Coda': '尾声，结尾',
    'Col legno': '用木头部分击打弦',
    'Con anima': '带有灵魂的，有激情的',
    'Con moto': '带有运动感的，带有速度的',
    'Crescendo': '渐强的',
    'D.C. al Fine': '从头开始，到终点为止',
    'dim.': '渐弱',
    'Diminuendo': '渐弱的',
    'Divisi': '分声部的',
    'Dolce': '甜美的，柔和的',
    'Doloroso': '悲伤的，哀伤的',
    'Energico': '充满能量的，强而有力的',
    'Espressivo': '富有表现力的',
    'Fermata': '停顿，暂停',
    'Forte': '强的',
    'Fortissimo': '非常强的',
    'Furioso': '愤怒的，狂暴的',
    'Giocoso': '欢快的，幽默的',
    'Glissando': '滑奏',
    'Grave': '庄板；慢而庄严的；严峻的',
    'Grazioso': '优雅的，柔美的',
    'Lamentoso': '悲叹的，哀伤的',
    'Largo': '广板；宽广的；庄严的',
    'loco': '回到标准音高(前面有8va)',
    'Legato': '连奏的',
    'Leggiero': '轻快的',
    'Lento': '慢板；慢慢的',
    'Lied': '歌曲（特别是德国艺术歌曲）',
    'Maestoso': '庄严的，威严的',
    'Movido': '活跃的',
    'Marcato': '强调的',
    'Marziale': '庄严的，进行曲风格的',
    'Meno': '少些；减少',
    'Meno mosso': '慢些（少快点）',
    'Misterioso': '神秘的',
    'molto accel.': '更快的加速',
    'molto rall.': '更多的渐慢以及扩张',
    'Non troppo': '不太，稍微',
    'Patetico': '悲怆的，感伤的',
    'Pesante': '慢速度；沉重的（每个音用一些重音）',
    'Pianissimo': '非常弱的',
    'Piano': '弱的',
    'Piu mosso': '更快些被激动了的',
    'Piu intenso': '更强烈的',
    'poco a poco': '逐渐地',
    'pizz': '把右手小指放在琴码处把弦蒙住，用其他指头拨弦，声音朦胧短促。',
    'Presto': '急板；迅速地',
    'Quasi': '类似于，好像',
    'Rallentando': '渐慢',
    'Risoluto': '坚定的，坚决的',
    'rit.': '渐慢的',
    'rall. poco a poco': '逐渐地渐慢',
    'Rubato': '自由变化的节奏',
    'Scherzando': '玩笑般的',
    'sempre stesso tempo': '始终保持相同的速度',
    'Scherzo': '玩笑曲',
    'Semplice': '简单的',
    'Serioso': '严肃的，认真地',
    'Sordino': '弱音器',
    'Sostenuto': '持续的，延长的',
    'Staccato': '断奏的',
    'stacc. e marcato': '断奏且强调的',
    'Stretto': '加速紧凑些',
    'Subito': '突然(sub.)',
    'Tenuto': '保持音符的时值',
    'Tambora':'像打击乐器一样把手放在音对应的弦上并移到靠近琴码的位置，使用i m a 指依次敲击弦',
    'Tranquillo': '平静的',
    'Tempo primo': '原速，速度还原全曲速度',
    'Tempo primo - deciso': '回到原速且坚定地演奏',
    'Tutti': '全体，所有乐器',
    'Vivace': '活跃的；快速的；敏捷的',
    'Vivo': '活泼；生动地；充满活力的',
    'Volante': '飞快的，轻盈的',
    'Zart': '柔和的，优美的',
    'Ab libitum': '随意的，自由发挥的',
    'Bellicoso': '战斗般的，好战的',
    'Bravura': '华丽的，炫技的',
    'Calando': '渐弱且渐慢',
    'Capriccioso': '随性的，任性的',
    'Con spirito': '带有精神的',
    'Con tenerezza': '温柔地',
    'Deciso': '坚定的，果断的',
    'Delicato': '精致的，优美的',
    'Drammatico': '戏剧性的',
    'Elegante': '优雅的，高贵的',
    'Eroico': '英雄的，壮丽的',
    'Feroce': '凶猛的，激烈的',
    'Festivo': '欢庆的，节日般的',
    'Fluente': '流畅的',
    'Grandioso': '宏伟的，壮丽的',
    'Grazioso con moto': '优美且富有动感的',
    'Impetuoso': '急促的，热情的',
    'Incalzando': '加急渐快并渐强',
    'Inquieto': '不安的，焦躁的',
    'Lacrimoso': '悲伤的，哭泣般的',
    'Languido': '柔弱的，无力的',
    'Liberamente': '自由的',
    'Maestoso con brio': '庄严且活力充沛的',
    'Melancolico': '忧郁的',
    'Misterioso con moto': '神秘且带有动感的',
    'Nobilmente': '高贵的，庄重的',
    'Nostalgico': '怀旧的',
    'Ondeggiante': '波动的，起伏的',
    'Ostinato': '固执的，反复的',
    'Pastorale': '田园风格的',
    'Patetico con dolore': '悲怆且痛苦的',
    'Pomposo': '浮夸的，华丽的',
    'Religioso': '虔诚的，宗教般的',
    'Rigoroso': '严格的，精确的',
    'Rustico': '乡村风格的',
    'Sognando': '梦幻般的',
    'Solenne': '庄严的',
    'Spianato': '平静的，均匀的',
    'Spirituoso': '精神饱满的',
    'Tenebroso': '阴暗的，神秘的',
    'Tosto': '快速的，敏捷的',
    'Trionfale': '凯旋的，胜利的',
    'Un poco': '稍微，一点',
    'Unisono': '同音的',
    'Vigoroso': '有力的，充满活力的',
    'Virtuoso': '炫技的，精湛的',
    'Zeloso': '热情的，热忱的',
    'Zingarese': '吉普赛风格的',
    'Alzapua': '弗拉门戈吉他技巧，用拇指连续向下和向上扫弦',
    'Apoyando': '支撑式击弦，手指击弦后停在相邻的弦上',
    'Arrastre': '快速滑弦技巧，从低音向高音滑动',
    'Barré': '横按，用食指同时按压多根弦',
    'Battente': '用手掌击打琴面制造打击乐效果',
    'Campanella': '钟声效果，使用不同把位的空弦和按弦产生重叠音效',
    'Ceja': '横按记号，标记需要横按的把位',
    'Cross-string': '跨弦技巧，在不同弦上演奏旋律',
    'Demi-barré': '半横按，只按压部分弦',
    'Descending rasgueado': '向下扫弦，从高音弦向低音弦方向',
    'Double-stop': '双音技巧，同时按住两个音',
    'Etude': '练习曲，专门用于提高某项技巧的曲目',
    'Fan rasgueado': '扇形扫弦，依次使用多个手指扫弦',
    'Fingerpicking': '指弹技巧，使用手指而不是拨片演奏',
    'Golpe': '击打琴面或琴弦制造打击乐效果',
    'Hammer-on': '锤弦，在已经发声的弦上用力按压产生新音',
    'Harmonics': '泛音技巧，在特定位置轻触琴弦产生泛音',
    'Hinge-barré': '铰链式横按，使用食指第一关节横按',
    'Inverted mordent': '倒波音，主音-上方邻音-主音的快速演奏',
    'Ligado': '连音技巧，包括锤弦和勾弦',
    'Mordent': '波音，主音-下方邻音-主音的快速演奏',
    'Natural harmonics': '自然泛音，在琴弦特定分数点处产生的泛音',
    'Ornament': '装饰音，用于装饰主要旋律的附加音符',
    'Picado': '单音技巧，使用i、m、a指头演奏单个音符',
    'Ponticello': '靠近琴桥处演奏，产生明亮尖锐的音色',
    'Position': '把位，左手在指板上的位置',
    'Pull-off': '勾弦，通过向下勾起手指产生新音',
    'Rasgueado': '扫弦技巧，连续使用多个手指快速扫弦',
    'Slide': '滑弦，从一个音滑向另一个音',
    'Slur': '连音，两个或多个音符连续演奏不间断',
    'Snap pizzicato': '巴托克拨弦，用力拨弦使其撞击指板',
    'Sul ponticello': '靠近琴桥演奏，产生明亮的金属音色',
    'Sul tasto': '靠近指板演奏，产生柔和的音色',
    'Tapping': '双手敲击指板产生音符',
    'Technique': '演奏技巧，包括各种手指动作和发声方法',
    'Tirando': '空弦式击弦，手指击弦后不停在相邻的弦上',
    'Tremolo': '颤音技巧，快速重复演奏同一个音',
    'Trill': '震音，两个相邻音符的快速交替',
    'Vibrato': '(vib.)揉弦，通过左手手指的快速摇动改变音高',
    'vib.': '(Vibtato) 揉弦，通过左手手指的快速摇动改变音高',
    'Artificial harmonics': '人工泛音，通过特殊技巧产生的泛音',
    'Bass string': '低音弦，吉他的第4、5、6弦',
    'Bridge saddle': '琴桥上的骨条，支撑琴弦的部件',
    'Classical position': '古典吉他演奏姿势，使用脚凳抬高左腿',
    'Damping': '制音，使用手掌或手指停止琴弦振动',
    'Extended technique': '扩展技巧，非传统的演奏方法',
    'Fret marker': '品位标记，指板上的位置标记',
    'Guitar support': '吉他支撑架，替代传统脚凳的现代设备',
    'Half-barré': '半横按，只按压部分弦的横按技巧',
    'Index finger': '食指，右手演奏中用i表示',
    'Jump position': '跳把位，快速改变左手位置',
    'Key position': '把位，左手在指板上的位置标记',
    'Left hand position': '左手位置，指左手在指板上的放置位置',
    'Middle finger': '中指，右手演奏中用m表示',
    'Nail shape': '指甲形状，古典吉他演奏者的指甲修剪方式',
    'Open string': '空弦，不按压时的琴弦音',
    'Palm mute': '手掌制音，用手掌轻触琴弦',
    'Quick position': '快速把位转换技巧',
    'Ring finger': '无名指，右手演奏中用a表示',
    'String crossing': '跨弦技巧，在不同弦间移动',
    'Thumb position': '拇指位置，右手拇指的放置位置',
    'Upper position': '高把位，在指板较高位置的演奏',
    'Warm tone': '温暖音色，靠近指板处演奏的音色特征',
    'Xylophone effect': '木琴效果，特殊的演奏技巧模仿木琴音色',
    'Yield point': '让音点，装饰音的结束点',
    'Zero fret': '零品，某些吉他在琴枕位置的金属品',
    'Aguado position': '阿瓜多演奏姿势，一种传统的古典吉他演奏姿势',
    'Barré chord': '横按和弦，使用横按技巧按压的和弦',
    'Capo position': '变调夹位置，变调夹放置的位置',
    'Dolce tone': '甜美音色，一种柔和优美的音色效果',
    'Ergonomic position': '人体工学演奏姿势，现代改良的演奏姿势',
    'Finger independence': '手指独立性，各手指独立运动的能力',
    'Hand angle': '手部角度，左右手相对于琴弦的角度',
    'Intimate tone': '私密音色，极其柔和的演奏音色',
    'Joint flexibility': '关节灵活性，手指关节的活动能力',
    'Knuckle position': '指关节位置，手指关节的放置位置',
    'Lateral movement': '横向移动，左手在指板上的横向移动',
    'Melodic line': '旋律线条，主要旋律的进行方式',
    'Nail angle': '指甲角度，击弦时指甲与弦的角度',
    'Octave harmonic': '八度泛音，在八度位置产生的泛音',
    'Parallel movement': '平行移动，和弦的平行移动技巧',
    'Quick release': '快速放松，击弦后的快速放松动作',
    'Relaxed position': '放松姿势，保持身体放松的演奏姿势',
    'String height': '弦高，琴弦距离指板的高度',
    'Thumb stroke': '拇指击弦，使用拇指演奏的技巧',
    'Upper neighbor': '上邻音，装饰音中高于主音的音',
    'Vertical movement': '垂直移动，手指的上下运动',
    'Wrist position': '手腕位置，右手手腕的放置位置',
    'Alternating bass': '交替低音，在伴奏中交替使用不同的低音',
    'Block chord': '块状和弦，同时按下所有和弦音',
    'Chord melody': '和弦旋律，将旋律和和弦结合的演奏方式',
    'Diagonal movement': '斜向移动，手指在指板上的斜向移动',
    'Expression mark': '表情记号，乐谱中的表现力指示',
    'Finger pattern': '指法模式，特定乐段的固定指法',
    'Grace note': '装饰音，快速演奏的装饰性音符',
    'Hand position': '手位，左右手的基本放置位置',
    'Inner voice': '内声部，多声部音乐中的中间声部',
    'Key change': '转调，音乐中的调式变化',
    'Lower neighbor': '下邻音，装饰音中低于主音的音',
    'Melodic minor': '旋律小调，上行和下行使用不同音阶',
    'Note duration': '音符时值，音符持续的时间长度',
    'Outer voice': '外声部，多声部音乐中的最高或最低声部',
    'Passing tone': '经过音，连接两个主要音符的音',
    'Quarter tone': '四分之一音，比半音还要小的音程',
    'Rest position': '休止位置，演奏间歇时手的放置位置',
    'Tension release': '张弛关系，音乐中的紧张与放松',
    'Unison note': '同音，相同音高的两个音',
    'Voice balance': '声部平衡，多声部间的音量平衡',
    'Walking bass': '行走低音，循序渐进的低音进行',
    'Altered chord': '变化和弦，包含变化音的和弦',
    'Broken chord': '分解和弦，将和弦音依次演奏',
    'Chromatic run': '半音进行，连续的半音音阶片段',
    'Double stop': '双音，同时演奏两个音的技巧',
    'Echo effect': '回声效果，模仿回声的演奏效果',
    'Guide tone': '引导音，和声进行中的关键音',
    'Harmonic minor': '和声小调，小调音阶的一种变体',
    'Implied harmony': '隐含和声，单音旋律暗示的和声',
    'Jazz chord': '爵士和弦，爵士音乐中常用的和弦',
    'Key signature': '调号，乐谱开始处的调式标记',
    'Leading tone': '导音，音阶中的第七音',
    'Modal mixture': '调式混合，大小调色彩的交替使用',
    'Neighbor tone': '邻音，装饰主音的上下相邻音',
    'Ostinato bass': '固定低音，重复的低音模式',
    'Parallel fifth': '平行五度，声部进行中的平行五度关系',
    'Quartal harmony': '四度和声，以四度叠置构成的和声',
    'Relative minor': '关系小调，与大调相关联的小调',
    'Secondary dominant': '副属和弦，非主调的属和弦',
    'Tertian harmony': '三度和声，以三度叠置构成的和声',
    'Unison passage': '齐奏段落，多声部同音进行的段落',
    'Voice crossing': '声部交叉，不同声部音高的交叉关系',
    'Whole tone': '全音，两个相邻全音阶音之间的音程',
    'Added tone': '附加音，和弦中添加的非和弦音',
    'Bridge passage': '过渡段，连接两个主要段落的部分',
    'Cadential': '终止式的，与乐句结束相关的',
    'Diatonic scale': '自然音阶，不含变化音的音阶',
    'Enharmonic note': '异名同音，写法不同但音高相同的音',
    'Final cadence': '最终终止，乐曲结束处的终止式',
    'Ground bass': '固定低音，不断重复的低音音型',
    'Harmonic rhythm': '和声节奏，和声变化的规律',
    'Interval study': '音程练习，针对特定音程的练习',
    'Just intonation': '纯律，基于自然泛音列的调律法',
    'Key center': '调中心，音乐的调性中心',
    'Linear motion': '线性进行，声部的水平移动',
    'Metric accent': '节拍重音，基于节奏的重音',
    'Non-chord tone': '非和弦音，不属于和弦的音',
    'Oblique motion': '斜行，一个声部保持不动而另一个声部移动',
    'Phrase mark': '乐句记号，标示乐句的记号',
    'Quality change': '音质变化，和弦或音程性质的改变',
    'Rhythmic motive': '节奏动机，重复出现的节奏型',
    'Scale degree': '音阶级数，音阶中音的位置序号',
    'Tonal center': '调性中心，音乐的主要调性依归',
    'Unstable tone': '不稳定音，倾向于解决到稳定音的音',
    'Voice leading': '声部进行，多声部音乐中各声部的移动方式',
    'Whole step': '全音，相当于两个半音的音程距离',
    'Altered scale': '变化音阶，包含多个变化音的音阶',
    'Borrowed chord': '借用和弦，从其他调式借用的和弦',
    'Chord symbol': '和弦记号，表示和弦的符号',
    'Dominant prep': '属和弦准备，为属和弦做准备的和弦',
    'Embellishment': '装饰，对基本音符的装饰',
    'Figured bass': '数字低音，巴洛克时期的和声记号',
    'Guide finger': '引导指，换把位时的参考手指',
    'Harmonic series': '泛音列，自然泛音的序列',
    'Inner melody': '内部旋律，和声中的中间声部旋律',
    'Jazz pattern': '爵士句型，典型的爵士乐句模式',
    'Key relation': '调式关系，不同调式间的关系',
    'Lydian mode': '利底亚调式，教会调式之一',
    'Modal chord': '调式和弦，基于特定调式的和弦',
    'Neapolitan': '那不勒斯和弦，小调式中的特殊和弦',
    'Octave study': '八度练习，针对八度音程的练习',
    'Parallel motion': '平行进行，声部同向平行移动',
    'Quintal harmony': '五度和声，以五度叠置构成的和声',
    'Root position': '原位，和弦的基本位置',
    'Scale pattern': '音阶模式，特定的音阶指法模式',
    'Tension chord': '紧张和弦，具有不稳定性的和弦',
    'Upper structure': '上部结构，和弦的上部音组合',
    'Voice exchange': '声部交换，声部之间的位置互换',
    'Whole tone scale': '全音音阶，由全音构成的音阶',
    'Augmented sixth': '增六和弦，包含增六度的特殊和弦',
    'Borrowed mode': '借用调式，从其他调式借用的音阶',
    'Chord voicing': '和弦排列，和弦音的具体排列方式',
    'Diminished scale': '减音阶，由半音和全音交替构成的音阶',
    'Enharmonic chord': '异名同音和弦，写法不同但音高相同的和弦',
    'False relation': '假关系，不同声部间的不协和关系',
    'Guide tone line': '引导音线，和声进行中的关键音进行',
    'Harmonic analysis': '和声分析，分析音乐的和声结构',
    'Inner voice leading': '内声部进行，中间声部的移动方式',
    'Jazz voicing': '爵士和弦排列，爵士风格的和弦排列',
    'Key modulation': '转调，从一个调式转到另一个调式',
    'Linear harmony': '线性和声，基于声部进行的和声',
    'Modal interchange': '调式交替，不同调式的交替使用',
    'Natural minor': '自然小调，小调的基本形式',
    'Octave displacement': '八度移位，将音符移动八度',
    'Parallel harmony': '平行和声，和声的平行移动',
    'Quartal voicing': '四度和声排列，以四度叠置的和声排列',
    'Root movement': '根音进行，和弦根音的移动方式',
    'Scale harmonization': '音阶和声化，为音阶配置和声',
    'Upper neighbor tone': '上邻音，高于主音的装饰音',
    'Voice separation': '声部分离，保持声部间的独立性',
    'Whole step motion': '全音进行，以全音距离移动的进行方式',
    'Accelerado': '渐快',
    'ad libitum': '速度任意',
    'Aentoq': '缓',
    'Affrettando': '再快一些',
    'Allargando': '逐渐变得宽广并渐强（通常用在曲尾）',
    'Allegro Assai': '很快的快板',
    'Allegro Moderato': '中庸的快板',
    'Allegro vivace': '活泼的快板',
    'Andamtte': '行板',
    'Andantino': '小行板；比Andanto稍快；较快的行板',
    'Dcppio movimento': '加快一倍',
    'Larghetto': '小广板；比Largo稍快；较快的小广板',
    'Moderatamente': '中等的；中庸的',
    'Moderato': '中板；适中；节制的',
    'Piu': '更；更多的',
    'Piu allegro': '更多的加快突快',
    'Piu lento': '更慢的',
    'Precipitano': '突然加快匆忙的急促的',
    'Presser': '匆忙；"赶"',
    'Prestissimo': '狂板；最快；极急速的临时转换速度',
    'Ritardato': '放慢了的',
    'Stringendo': '加紧加快',
    'Vivacissimo': '最急板；十分活跃地；非常爽快地',
    'Accel': '渐快',
    'Con forza': '有力的',
    'Con tutta forza': '用全部力量，尽可能响亮地',
    'Declamando': '朗诵般的',
    'Decrescendo': '渐弱',
    'Deliberato': '果断的',
    'Dolente': '悲哀地，怨诉地',
    'Elegance': '优雅的，风流潇洒的',
    'Elegiaco': '哀悼的，挽歌的',
    'Elevato': '崇高的，升华的',
    'Fastoso': '显赫辉煌地，华美地',
    'Fiaccamente': '柔弱的，无力的',
    'Fiero': '骄傲的，激烈的',
    'Flessibile': '柔和而可伸缩的',
    'Fortepiano': '强后突弱',
    'Frescamente': '新鲜地，清新地，年轻地',
    'Fuocoso': '火热的，热烈的',
    'Generoso': '慷慨的，宽宏而高贵的',
    'Gioiante': '快乐的',
    'Giusto': '准确的，适当的',
    'Imitando': '仿拟，模仿',
    'Incalcando': '加急渐快并渐强',
    'Indifferenza': '冷漠',
    'Infantile': '幼稚的，孩子气的',
    'Innig': '内在的，深切的',
    'Innocente': '天真无邪的，坦率无知的，纯朴的',
    'Inquielo': '不安的',
    'Irato': '发怒地',
    'Largamente': '广阔的，浩大地',
    'Leggiere': '轻巧的，轻快的',
    'Luttuoso': '哀痛的',
    'Meno allegro': '速度转慢',
    'Mesto': '忧郁的',
    'Mezzo Forte': '中强',
    'Mezzo Piano': '中弱',
    'Mormorando': '絮絮低语地',
    'Netto': '清爽的，清楚的',
    'Nobile': '高贵的',
    'Pacatamente': '平静地，温和地，安详地',
    'Parlante': '说话似的',
    'Pianoforte': '先弱后强',
    'Piacevole': '愉快的，惬意的，使人悦愉的',
    'Placido': '平静的',
    'Ponderoso': '有力而印象深刻的',
    'Posato': '安祥的，稳重的',
    'Pressante': '紧迫的，催赶的',
    'Riposato': '恬静的',
    'Ritardando': '渐慢',
    'Sciolto': '无拘无束的，流畅的，敏捷的',
    'Sdegnoso': '愤慨的，轻蔑的',
    'Sforzando': '突强',
    'Smania': '狂怒，疯狂',
    'Smorfioso': '做作的，买弄的',
    'Smorzando': '渐弱并渐慢、逐渐消失',
    'Soave': '柔和的，甘美的',
    'Solennemente': '庄严而隆重地',
    'Sonoro': '响亮的',
    'Sotto voce': '轻声的，弱声的',
    'Spirito': '精神饱满的，热情的，有兴致的，鼓舞的',
    'Strappando': '用粗暴的力量来发音',
    'Teneramente': '温柔地，柔情地',
    'Timoroso': '胆怯的，恐惧的，畏惧的',
    'Trionfante': '凯旋的',
    'Tumultuoso': '嘈杂的，喧闹的',
    'Veloce': '敏捷的',
    'Vezzoso': '优美的，可爱的',
    'Violentemente': '狂暴地，猛烈地',
    'Vivente': '活泼的，有生气的',
    'Vivido': '活跃的，栩栩如生的',
    'Zeffiroso': '轻盈柔和的，象微风般的',
    'Accento': '加强地，突强，特重',
    'Forzato': '加强地，突强，特重',
    'Rinforzando': '加强地，突强，特重',
    'Bitonalità': '双调性',
    'Barre': '横按，横压',
    'Capotasto': '变调夹',
    'Harmonic': '泛音',
    'Natural harmonic': '自然泛音',
    'Artificial harmonic': '人工泛音',
    'Pizzicato': '拨奏，闷音',
    'Bend': '推弦',
    'Release': '放弦',
    'Pre-bend': '预推弦',
    'Unison bend': '同音推弦',
    'Pinch harmonic': '人工泛音（挤压式）',
    'Sweep picking': '扫拨',
    'Economy picking': '经济拨片法',
    'Alternate picking': '交替拨片法',
    'String skipping': '跳弦',
    'Hybrid picking': '混合拨奏法',
    'Travis picking': '特拉维斯指法',
    'Artificial harmonic slide': '人工泛音滑音',
    'Legato slide': '连奏滑音',
    'Shift slide': '换把滑音',
    'Ghost note': '装饰音，幽灵音',

    # 巴洛克时期
    'Allemande': '德国舞曲（巴洛克组曲中的第一乐章）',
    'Aria da capo': '带反复的咏叹调',
    'Bourree': '布雷舞曲（法国快速舞曲）',
    'Cantus firmus': '定旋律（固定音调）',
    'Chaconne': '恰空舞曲（庄严的三拍子）',
    'Concerto grosso': '大协奏曲',
    'Courante': '库朗特舞曲（法国快速舞曲）',
    'Gavotte': '加沃特舞曲（法国舞曲）',
    'Gigue': '吉格舞曲（快速的复拍子舞曲）',
    'Menuet': '小步舞曲',
    'Passacaglia': '帕萨卡利亚（庄严的变奏曲）',
    'Pavane': '孔雀舞曲（庄重的舞曲）',
    'Ricercare': '里切卡雷（复调音乐）',
    'Sarabande': '萨拉班德舞曲（庄严的舞曲）',
    'Toccata': '托卡塔（炫技性乐曲）',

    # 古典主义时期
    'Cadenza': '华彩段',
    'Divertimento': '嬉游曲',
    'Eingang': '引入句（德语）',
    'Exposition': '呈示部',
    'Galant': '优雅的风格',
    'Hauptsatz': '主题（德语）',
    'Mannheimer Schule': '曼海姆乐派',
    'Rondo': '回旋曲',
    'Seitensatz': '副主题（德语）',
    'Serenata': '小夜曲',
    'Sinfonia': '交响曲',
    'Sonatina': '小奏鸣曲',
    'Sturm und Drang': '狂飙运动（德语）',

    # 浪漫主义时期
    'Albumblatt': '相册页（德语）',
    'Ballade': '叙事曲',
    'Berceuse': '摇篮曲（法语）',
    'Capriccio': '随想曲',
    'Charakterstück': '性格小品（德语）',
    'Consolation': '安慰曲（法语）',
    'Fantasie': '幻想曲（德语）',
    'Impromptu': '即兴曲（法语）',
    'Innigkeit': '内在的，深情的（德语）',
    'Intermezzo': '间奏曲',
    'Leidenschaftlich': '热情地（德语）',
    'Lied ohne Worte': '无词歌（德语）',
    'Märchen': '童话（德语）',
    'Moment musical': '音乐瞬间（法语）',
    'Nachtstück': '夜曲（德语）',
    'Nocturne': '夜曲（法语）',
    'Novellette': '小说曲',
    'Rhapsodie': '狂想曲',
    'Romanze': '浪漫曲（德语）',
    'Schwungvoll': '充满激情地（德语）',
    'Sehnsucht': '渴望（德语）',
    'Träumerei': '梦幻曲（德语）',
    'Waldszenen': '森林景象（德语）',
    'Zeitmaß': '速度（德语）',

    # 意大利语专业术语
    'A capriccio': '随意地',
    'A piacere': '随心所欲地',
    'Abbandono': '放任的',
    'Acciaccatura': '短倚音',
    'Affannato': '焦虑的',
    'Agilità': '灵活的',
    'Con abbandono': '放纵地',
    'Con alcuna licenza': '带有一些自由',
    'Con amore': '充满爱意地',
    'Con bravura': '技巧性地',
    'Con calore': '热情地',
    'Con delicatezza': '细腻地',
    'Con dolcezza': '甜美地',
    'Con duolo': '悲伤地',
    'Con eleganza': '优雅地',
    'Con espressione': '富有表情地',
    'Con fierezza': '自豪地',
    'Con fuoco': '火热地',

    # 法语专业术语
    'À deux mains': '双手的',
    'À plein son': '全音的',
    'Avec âme': '有灵魂地',
    'Avec chaleur': '热情地',
    'Avec douleur': '悲痛地',
    'Avec éclat': '辉煌地',
    'Avec élan': '热情奔放地',
    'Avec émotion': '激动地',
    'Avec grâce': '优美地',
    'Avec largeur': '宽广地',
    'Avec passion': '热情地',
    'Brillamment': '辉煌地',
    'Chantant': '如歌的',
    'Délicatement': '精致地',
    'Doucement': '温柔地',
    'En dehors': '突出地',
    'Gracieusement': '优雅地',
    'Légèrement': '轻盈地',
    'Passionnément': '热情地',
    'Sans rigueur': '自由地',

    # 德语专业术语
    'Ausdruck': '表情',
    'Ausdrucksvoll': '富有表情地',
    'Bewegt': '激动的',
    'Dämpfer': '弱音器',
    'Einfach': '简单地',
    'Empfindung': '感觉',
    'Entfernt': '远离的',
    'Entschieden': '坚决地',
    'Erhaben': '崇高的',
    'Feurig': '火热地',
    'Fließend': '流畅地',
    'Gebunden': '连奏的',
    'Gefühlvoll': '感情丰富地',
    'Gehalten': '持续的',
    'Gemächlich': '从容地',
    'Gesangvoll': '如歌的',
    'Getragen': '庄重地',
    'Gewichtig': '重要地',
    'Heftig': '激烈地',
}

# 添加更多术语
music_terms.update({
    'Accento': '加强地，突强，特重',
    'Calando': '渐弱，渐安静',
    'port.': '(portamento）在两个音符之间平滑地滑动或连音的技巧',
    'portamento': '(port.) 在两个音符之间平滑地滑动或连音的技巧',
    'plp.': '拨弦',
    'unghia': '指甲，在吉他等弦乐器的演奏中，"unghia" 通常指示演奏者使用指甲进行拨弦演奏。这种技巧可以产生明亮、清晰的音色，常用于古典吉他演奏中。例如，"unghia" 可能出现在乐谱上，指示演奏者使用指甲拨弦。此外，"unghia" 也可能与其他演奏技巧结合使用，如**"pizzicato"（拨奏）或"arco"（用弓拉奏），以指示演奏者在特定段落使用指甲拨弦。',
    'unghiata': '(ung.) 用指甲弹奏 拨弦',
    'ung.': '(unghiata) 用指甲弹奏 拨弦',
    'molto': '非常。很多',
    'unis.': '(unisono)齐奏，同声',
    'unisono': '(unis.) 齐奏，同声',
    'pred.': '(perdendosi/perdere) 渐渐消失 逐渐减弱',
    'Verso il ponticello': '朝向琴桥处，指在古典吉他演奏中，将右手的指尖或拨片靠近琴桥的位置拨弦。这样可以产生更加清晰、尖锐的音色，常用于强调或创造紧张的音效。',
    'Pesante': '(pesante) 重的、沉重的，通常指演奏时要用较大的力度或更加强调的方式，表现出一种沉稳、庄重或强烈的情感。在音乐中，"pesante" 通常要求演奏者在音量和力度上表现出沉重感。',
    'En dehors': '突出，指在演奏中某个声部或音符应特别突出或强。',
    'Crescendo': '渐强，指音量逐渐增大，通常会持续一段时间，直到达到一个较大的音量。',
    'Poco': '稍微，指音乐中某个动作或特征应当轻微或略微地进行，通常用于修饰其他术语，表示程度较轻',
    'A': '以、通过，通常用于连接其他术语，表示某个动作或特征的方向或方式。在音乐中，它可以指示演奏方式或手段。',
    'gliss.': '(glissando) 滑奏，指的是在两个音符之间平滑地滑动，通常是用弓、手指或其他工具，使音符之间没有明确的间隔。',
    'più': '更多，用于修饰其他术语，表示增加或更强烈的意思，常见于动态和表现力指示中，如 "più forte"（更响亮）。',
    'tast.': '(tastiera) 键盘，指的是钢琴或类似乐器的键盘部分。',
    'secco': '干的，指演奏时音符或和弦的音色应当简洁、干脆，通常意味着没有过多的延续或共鸣。',
    'lunga': '长的，指音符或节奏的延长，表示该音符或休止符应该持续较长时间。',
    'sempre': '始终，指音乐中某个动作或效果应持续进行，直到另有指示。',
    'poco vib.': '(poco vibrato) 稍微揉弦，指在演奏中使用轻微的颤音效果。',
    'tastiera': '键盘，指钢琴或其他具有键盘的乐器的按键部分。',
    'dolcissimo': '非常柔美，指以极其温柔、甜美的方式演奏。',
    'a tempo, poco stringendo': '按原速，稍微加速。指演奏时恢复到原来的节奏，并在接下来的部分稍微加快速度。',
    'Stringendo': '渐快，指逐渐加速，通常用于表示音乐应加快节奏。',
    'marcato': '强调，指在演奏中某个音符或节奏应特别突出和强烈，通常用较大的力度表现。',






    # ... [中间的术语保持不变]
})

# 添加吉他术语
music_terms.update({
    'Apoyando': '支撑式扫弦，休止式扫弦',
    'Arpeggio': '琶音，分散和弦',
    # ... [后面的术语保持不变]
})

# 去除助词的函数
def remove_particles(text):
    """去除中文助词（的、地、得）"""
    return text.replace('的', '').replace('地', '').replace('得', '')

# 在文件开头添加类定义
class MusicDictionary:
    def __init__(self):
        """初始化音乐术语字典"""
        self.music_terms = music_terms
        self.lower_to_original = {k.lower(): k for k in self.music_terms.keys()}
        
        # 创建隐藏的数据目录
        self.data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.music_dict_data')
        os.makedirs(self.data_dir, exist_ok=True)
        
        # 初始化收藏夹和历史记录文件（使用隐藏文件）
        self.favorites_file = os.path.join(self.data_dir, '.favorites.json')
        self.history_file = os.path.join(self.data_dir, '.history.json')
        
        # 加载收藏夹和搜索历史
        self.favorites = self._load_favorites()
        self.search_history = self.load_history()
        
        # 添加版本信息
        self.version = "1.0.0"
        
        # 检查更新（如果需要）
        self._check_updates()

    def _load_favorites(self):
        """加载收藏夹"""
        try:
            if os.path.exists(self.favorites_file):
                with open(self.favorites_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            print(f"加载收藏夹失败：{str(e)}")
            return {}

    def save_favorites(self):
        """保存收藏夹"""
        try:
            with open(self.favorites_file, 'w', encoding='utf-8') as f:
                json.dump(self.favorites, f, ensure_ascii=False, indent=4)
            return True
        except Exception as e:
            print(f"保存收藏夹失败：{str(e)}")
            return False

    def query_term(self, term):
        """查询术语"""
        # 处理特殊命令
        if term == '?':
            return self.show_help(), None, None
        
        # 清理输入（移除特殊字符）
        term = ''.join(c for c in term if c.isalnum() or c.isspace() or '\u4e00' <= c <= '\u9fff')
        
        if not term or term.isspace():
            return "请输入要查询的内容", None, None

        term_lower = term.lower()
        term_length = len(term_lower)
        matches = []
        similarity_scores = {}
        result = ""

        # 1. 精确匹配（不区分大小写）
        if term_lower in self.lower_to_original:
            original_key = self.lower_to_original[term_lower]
            return f"\n找到术语：{original_key}\n释义：{self.music_terms[original_key]}", original_key, None

        # 2. 计算相似度和匹配度
        term_pinyin = ''.join(lazy_pinyin(term))  # 搜索词的拼音
        potential_matches = []

        for original_term, meaning in self.music_terms.items():
            original_term_lower = original_term.lower()
            score = 0
            match_type = set()

            # 中文匹配（最高优先级）
            if term in original_term or term in meaning:
                score += 1.0  # 提高中文直接匹配的权重
                match_type.add('chinese')
                
            # 拼音匹配
            original_pinyin = ''.join(lazy_pinyin(original_term))  # 术语的拼音
            meaning_pinyin = ''.join(lazy_pinyin(meaning))  # 释义的拼音
            
            # 完整拼音匹配
            if term_pinyin == original_pinyin or term_pinyin in meaning_pinyin.split():
                score += 0.9
                match_type.add('pinyin_full')
            # 部分拼音匹配
            elif term_pinyin in original_pinyin or term_pinyin in meaning_pinyin:
                score += 0.7
                match_type.add('pinyin_partial')

            # 英文术语的特殊处理
            if all(ord(c) < 128 for c in original_term):
                # 前缀匹配
                if original_term_lower.startswith(term_lower):
                    score += 0.8
                    match_type.add('prefix')
                # 包含匹配
                elif term_lower in original_term_lower:
                    pos = original_term_lower.index(term_lower)
                    if pos == 0:
                        score += 0.7
                    else:
                        score += 0.5
                    match_type.add('contain')
                
                # 编辑距离相似度
                similarity = difflib.SequenceMatcher(None, term_lower, original_term_lower).ratio()
                if similarity > 0.6:
                    score += similarity * 0.4
                    match_type.add('similar')

            # 确保总分不超过1.0
            final_score = min(score, 1.0)
            
            # 提高匹配阈值
            if final_score > 0.3:  # 只保留相关度较高的匹配
                potential_matches.append({
                    'term': original_term,
                    'score': final_score,
                    'match_type': match_type,
                    'meaning': meaning,
                    'term_length': len(original_term)  # 添加长度信息用于排序
                })

        # 优化排序逻辑
        def sort_key(match):
            return (
                -match['score'],  # 首先按分数降序
                'chinese' in match['match_type'],  # 中文匹配优先
                'pinyin_full' in match['match_type'],  # 完整拼音匹配次之
                abs(match['term_length'] - term_length),  # 长度相近的优先
                match['term'].lower()  # 最后按字母顺序
            )

        # 排序并取前10个最佳匹配
        potential_matches.sort(key=sort_key)
        top_matches = potential_matches[:10]

        if top_matches:
            result = "\n找到相关术语：\n"
            for i, match in enumerate(top_matches, 1):
                score = int(match['score'] * 100)
                result += f"{i}. {match['term']}（匹配度：{score}%）\n   {match['meaning']}\n"
                matches.append(match['term'])
                similarity_scores[match['term']] = match['score']

        if len(matches) == 1:
            return result, matches[0], matches
        return result, None, matches

        # 优化未找到匹配时的提示信息
        if not potential_matches:
            suggestions = []
            if any('\u4e00' <= c <= '\u9fff' for c in term):  # 中文搜索
                suggestions.append("1. 尝试使用拼音搜索")
                suggestions.append("2. 尝试使用部分关键词")
                suggestions.append("3. 输入 ? 查看帮助信息")
            elif term.isalpha():  # 英文/拼音搜索
                suggestions.append("1. 检查拼写是否正确")
                suggestions.append("2. 尝试使用更短的关键词")
                suggestions.append("3. 尝试使用中文搜索")
            else:  # 其他情况
                suggestions.append("1. 请使用中文、英文或拼音进行搜索")
                suggestions.append("2. 避免使用特殊字符")
                suggestions.append("3. 输入 ? 查看帮助信息")
            
            return f"\n未找到与 '{term}' 相关的术语。\n\n建议：\n" + "\n".join(suggestions), None, None

    def handle_favorites_choice(self, matches):
        """处理多个匹配结果的收藏选择"""
        # 删除重复的结果显示，因为在query_term中已经显示过了
        while True:
            try:
                choice = input("\n请输入要收藏的术语编号（0取消）：").strip()
                if not choice:  # 处理空输入
                    print("请输入数字！")
                    continue
                    
                choice = int(choice)
                if choice == 0:
                    return
                if 1 <= choice <= len(matches):
                    selected_term = matches[choice - 1]
                    if selected_term in self.favorites:
                        print(f"\n'{selected_term}' 已在收藏夹中")
                        return
                    self.favorites[selected_term] = self.music_terms[selected_term]
                    self.save_favorites()
                    print(f"\n已收藏：{selected_term}")
                    return
                print("无效的编号，请重新输入！")
            except ValueError:
                print("请输入有效的数字！")

    def export_favorites(self):
        """导出收藏夹"""
        if not self.favorites:
            print("\n收藏夹为空，无法导出\n")
            return
            
        try:
            # 选择导出格式
            print("\n请选择导出格式：")
            print("1. 文本文件 (.txt)")
            print("2. Word文档 (.docx)")
            print("3. Excel表格 (.xlsx)")
            print("4. HTML网页 (.html)")
            
            while True:
                format_choice = input("\n请输入格式编号(1-4)：").strip()
                if format_choice in ['1', '2', '3', '4']:
                    break
                print("无效的选择，请重新输入！")

            # 选择导出位置
            print("\n请选择导出位置：")
            print("1. 桌面")
            print("2. 下载")
            print("3. 文档")
            print("4. 当前目录")
            
            while True:
                location = input("\n请输入数字(1-4)：").strip()
                if location in ['1', '2', '3', '4']:
                    break
                print("无效的选择，请重新输入！")
            
            # 设置导出路径
            export_paths = {
                '1': os.path.expanduser('~/Desktop'),
                '2': os.path.expanduser('~/Downloads'),
                '3': os.path.expanduser('~/Documents'),
                '4': os.path.dirname(os.path.abspath(__file__))
            }
            
            base_path = export_paths[location]
            if not os.path.exists(base_path):
                print(f"\n导出路径不存在：{base_path}")
                return

            # 根据选择的格式导出
            if format_choice == '1':  # 文本文件
                filename = os.path.join(base_path, 'music_terms_favorites.txt')
                self._export_as_text(filename)
                
            elif format_choice == '2':  # Word
                try:
                    from docx import Document
                    filename = os.path.join(base_path, 'music_terms_favorites.docx')
                    self._export_as_word(filename)
                except ImportError:
                    print("\n导出Word文档需要安装python-docx库")
                    print("请运行: pip install python-docx")
                    return
                    
            elif format_choice == '3':  # Excel
                try:
                    import pandas as pd
                    filename = os.path.join(base_path, 'music_terms_favorites.xlsx')
                    self._export_as_excel(filename)
                except ImportError:
                    print("\n导出Excel需要安装pandas库")
                    print("请运行: pip install pandas openpyxl")
                    return
                    
            elif format_choice == '4':  # HTML
                filename = os.path.join(base_path, 'music_terms_favorites.html')
                self._export_as_html(filename)

            print(f"\n收藏夹已成功导出到：{filename}\n")
            
        except Exception as e:
            print(f"\n导出失败：{str(e)}")
            print("请检查文件权限或选择其他导出位置")

    def _export_as_text(self, filename):
        """导出为文本文件"""
        if os.path.exists(filename):
            if input("\n文件已存在，是否覆盖？(y/n): ").strip().lower() != 'y':
                print("\n导出已取消")
                return False
            
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write("音乐术语词典 - 收藏夹\n")
                f.write("=" * 50 + "\n\n")
                for term, meaning in sorted(self.favorites.items()):
                    f.write(f"{term}:\n    {meaning}\n\n")
            
            # 验证导出的文件是否可以正确读取
            with open(filename, 'r', encoding='utf-8') as f:
                content = f.read()
                if not content:
                    raise Exception("导出文件为空")
            return True
        except UnicodeEncodeError:
            print("\n导出失败：编码错误，请检查文件名或内容是否包含特殊字符")
            return False
        except Exception as e:
            print(f"\n导出失败：{str(e)}")
            return False

    def _export_as_word(self, filename):
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            if os.path.exists(filename):
                if input("\n文件已存在，是否覆盖？(y/n): ").strip().lower() != 'y':
                    print("\n导出已取消")
                    return False
            
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('音乐术语词典 - 收藏夹', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for term, meaning in sorted(self.favorites.items()):
                # 添加术语
                p = doc.add_paragraph()
                run = p.add_run(term)
                run.bold = True
                run.font.size = Pt(14)
                
                # 添加释义
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(20)
                run = p.add_run(meaning)
                run.font.size = Pt(12)
                
                doc.add_paragraph()  # 添加空行
            
            doc.save(filename)
            return True
            
        except ImportError:
            print("\n导出Word文档需要安装python-docx库")
            print("请运行: pip install python-docx")
            return False
        except Exception as e:
            print(f"\n导出Word文档失败：{str(e)}")
            return False

    def _export_as_excel(self, filename):
        """导出为Excel表格"""
        try:
            import pandas as pd
            from openpyxl import load_workbook
            from openpyxl.styles import Alignment, Font
            
            if os.path.exists(filename):
                if input("\n文件已存在，是否覆盖？(y/n): ").strip().lower() != 'y':
                    print("\n导出已取消")
                    return False  # 直接返回，不执行后续的导出操作
            
            # 先用pandas创建基础Excel文件    
            df = pd.DataFrame(list(self.favorites.items()), columns=['术语', '释义'])
            df.to_excel(filename, index=False, sheet_name='收藏夹')
            
            # 然后用openpyxl优化格式
            wb = load_workbook(filename)
            ws = wb.active
            
            # 设置列宽
            ws.column_dimensions['A'].width = 30  # 术语列宽度
            ws.column_dimensions['B'].width = 80  # 释义列宽度
            
            # 设置表头格式
            header_font = Font(bold=True, size=12)
            for cell in ws[1]:
                cell.font = header_font
            
            # 设置所有单元格格式
            for row in ws.iter_rows(min_row=1):
                for cell in row:
                    # 启用自动换行和垂直居中对齐
                    cell.alignment = Alignment(
                        wrap_text=True,
                        vertical='center',
                        horizontal='left'
                    )
            
            # 自动调整行高
            for row in ws.iter_rows(min_row=2):  # 从第二行开始（跳过表头）
                max_lines = 1
                for cell in row:
                    if cell.value:  # 如果单元格有内容
                        # 计算文本换行后的行数
                        lines = str(cell.value).count('\n') + 1
                        # 考虑单元格宽度导致的自动换行
                        text_width = len(str(cell.value)) * 2  # 估计中文字符宽度
                        col_width = ws.column_dimensions[cell.column_letter].width
                        wrapped_lines = (text_width / col_width) + 1
                        max_lines = max(max_lines, lines, int(wrapped_lines))
                
                # 设置行高（每行文字大约需要20个单位）
                row_height = max_lines * 20
                ws.row_dimensions[row[0].row].height = row_height
            
            # 保存修改后的文件
            wb.save(filename)
            print(f"\n已成功导出到：{filename}")
            print("注意：已优化长文本显示格式")
            return True
            
        except ImportError:
            print("\n导出Excel需要安装pandas和openpyxl库")
            print("请运行: pip install pandas openpyxl")
            return False
        except Exception as e:
            print(f"\n导出Excel失败：{str(e)}")
            print("请检查文件是否被其他程序占用")
            return False

    def _export_as_html(self, filename):
        """导出为HTML网页"""
        if os.path.exists(filename):
            if input("\n文件已存在，是否覆盖？(y/n): ").strip().lower() != 'y':
                print("\n导出已取消")
                return False
            
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>音乐术语词典 - 收藏夹</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; }
                h1 { color: #333; text-align: center; }
                .term { color: #2c5282; font-size: 1.2em; margin-top: 20px; }
                .meaning { margin-left: 20px; color: #444; }
            </style>
        </head>
        <body>
            <h1>音乐术语词典 - 收藏夹</h1>
        """
        
        for term, meaning in sorted(self.favorites.items()):
            html_content += f"""
            <div class="term">{term}</div>
            <div class="meaning">{meaning}</div>
            """
        
        html_content += """
        </body>
        </html>
        """
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return True

    def load_history(self):
        """加载搜索历史"""
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return []
        except Exception as e:
            print(f"加载搜索历史失败：{str(e)}")
            return []

    def save_history(self):
        """保存搜索历史"""
        try:
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(self.search_history, f, ensure_ascii=False, indent=4)
            return True
        except Exception as e:
            print(f"保存搜索历史失败：{str(e)}")
            return False

    def add_to_history(self, term):
        """添加搜索记录"""
        if term in self.search_history:
            # 如果已存在，先移除旧记录
            self.search_history.remove(term)
        
        # 添加到历史记录开头
        self.search_history.insert(0, term)
        
        # 保持最多20条记录
        if len(self.search_history) > 20:
            self.search_history.pop()
        
        # 保存到文件
        self.save_history()

    def show_history(self):
        """显示搜索历史"""
        if not self.search_history:
            print("\n暂无搜索历史\n")
            return None
            
        print("\n最近的搜索记录：")
        for i, term in enumerate(self.search_history, 1):
            # 如果这个术语在收藏夹中，标记出来
            if term in self.favorites:
                print(f"{i}. {term} [已收藏]")
            else:
                print(f"{i}. {term}")
            
        print("\n提示：输入序号可以重新搜索对应术语")
        
        while True:
            choice = input("请输入序号(直接回车返回)：").strip()
            if not choice:  # 空输入，返回主菜单
                print()
                return None
                
            if choice.isdigit():
                idx = int(choice) - 1
                if 0 <= idx < len(self.search_history):
                    return self.search_history[idx]  # 返回选中的术语供重新搜索
                else:
                    print(f"请输入1-{len(self.search_history)}之间的数字！")
            else:
                print("请输入有效的数字！")

    def favorites_quiz(self):
        """收藏夹术语测验模式"""
        if not self.favorites:
            print("\n收藏夹为空，无法开始测验！请先添加一些术语到收藏夹。")
            return
        
        # 让用户选择测验数量
        total_terms = len(self.favorites)
        print(f"\n当前收藏夹共有 {total_terms} 个术语")
        while True:
            try:
                num = input(f"请输入要测试的题目数量(1-{total_terms}，直接回车默认5题)：").strip()
                if not num:  # 使用默认值
                    total = min(5, total_terms)
                    break
                num = int(num)
                if 1 <= num <= total_terms:
                    total = num
                    break
                print(f"请输入1到{total_terms}之间的数字！")
            except ValueError:
                print("请输入有效的数字！")
        
        score = 0
        wrong_terms = []  # 记录答错的题目
        
        print(f"\n开始收藏夹术语测验！共 {total} 题")
        # 从收藏夹中随机选择术语
        terms = random.sample(list(self.favorites.keys()), total)
        
        for i, term in enumerate(terms, 1):
            print(f"\n第{i}/{total}题：{term} 的含义是？")
            
            # 生成选项
            correct_answer = self.favorites[term]
            # 从其他术语中随机选择3个错误答案
            other_terms = [v for k, v in self.favorites.items() if k != term]
            if len(other_terms) >= 3:
                wrong_answers = random.sample(other_terms, 3)
            else:
                # 如果收藏夹术语不足4个，从主字典补充错误答案
                remaining_terms = [v for k, v in self.music_terms.items() 
                                 if k != term and v != correct_answer]
                wrong_answers = random.sample(remaining_terms, 3)
            
            # 将正确答案和错误答案混合
            all_answers = wrong_answers + [correct_answer]
            random.shuffle(all_answers)
            correct_index = all_answers.index(correct_answer)
            
            # 显示选项
            for j, answer in enumerate(all_answers, 1):
                print(f"{j}. {answer}")
            
            # 获取用户答案
            while True:
                try:
                    user_choice = input("\n请选择正确答案的编号(1-4，q退出测验)：").strip().lower()
                    if user_choice == 'q':
                        print("\n测验已终止")
                        return
                    user_choice = int(user_choice)
                    if 1 <= user_choice <= 4:
                        break
                    print("请输入1-4之间的数字！")
                except ValueError:
                    print("请输入有效的数字！")
            
            # 检查答案
            if all_answers[user_choice - 1] == correct_answer:
                print("✓ 回答正确！")
                score += 1
            else:
                print(f"✗ 回答错误。正确答案是：{correct_answer}")
                wrong_terms.append((term, correct_answer))  # 记录错误的题目
        
        # 显示测验结果
        print(f"\n测验结束！你的得分是：{score}/{total} ({score/total*100:.1f}%)")
        
        # 如果有答错的题目，显示错题回顾
        if wrong_terms:
            print("\n以下是答错的题目，建议重点复习：")
            for term, answer in wrong_terms:
                print(f"- {term}: {answer}")

    def show_favorites(self):
        """显示收藏夹内容"""
        if not self.favorites:
            print("\n收藏夹为空\n")
            return
            
        print("\n收藏夹内容：")
        for i, (term, meaning) in enumerate(self.favorites.items(), 1):
            print(f"{i}. {term}: {meaning}")
        print(f"\n共 {len(self.favorites)} 个收藏项")
        print(f"收藏夹保存位置：{self.favorites_file}")
        print("\n提示：输入rm移除单个术语，输入rm*移除所有术语")
        
        while True:
            choice = input().strip().lower()
            if not choice:  # 直接回车返回
                break
            elif choice == 'rm':
                self.remove_single_favorite()
                break
            elif choice == 'rm*':
                self.remove_all_favorites()
                break
        print()  # 空行分隔

    def remove_single_favorite(self):
        """移除单个收藏项"""
        if not self.favorites:
            print("\n收藏夹为空\n")
            return
            
        while True:
            try:
                choice = input(f"\n请输入要删除的术语编号（1-{len(self.favorites)}），输入0取消：").strip()
                if not choice.isdigit():
                    print("请输入有效的数字")
                    continue
                    
                idx = int(choice)
                if idx == 0:
                    return
                    
                if 1 <= idx <= len(self.favorites):
                    term_to_remove = list(self.favorites.keys())[idx-1]
                    del self.favorites[term_to_remove]
                    self.save_favorites()
                    print(f"已删除：{term_to_remove}")
                    # 重新显示收藏夹
                    self.show_favorites()
                    break
                else:
                    print(f"请输入1-{len(self.favorites)}之间的数字")
            except ValueError:
                print("请输入有效的数字")

    def remove_all_favorites(self):
        """移除所有收藏项"""
        if not self.favorites:
            print("\n收藏夹为空\n")
            return
            
        confirm = input("\n确定要删除所有收藏吗？此操作不可恢复(y/n): ").strip().lower()
        if confirm == 'y':
            self.favorites.clear()
            self.save_favorites()
            print("已清空所有收藏")
        else:
            print("已取消删除操作")

    def handle_search_result(self, original_term, matches):
        """处理搜索结果"""
        if original_term:  # 精确匹配的情况
            while True:
                choice = input("\n是否收藏该术语？(y/n): ").strip().lower()
                if choice == 'y':
                    if original_term in self.favorites:
                        print(f"\n'{original_term}' 已在收藏夹中")
                        return None
                    self.favorites[original_term] = self.music_terms[original_term]
                    if self.save_favorites():
                        print("收藏成功！\n")
                    else:
                        print("收藏失败，请检查文件权限\n")
                    return None
                elif choice == 'n':
                    return None
                else:
                    return choice  # 将非y/n的输入作为新的搜索词
        elif matches:  # 模糊匹配的情况
            while True:
                choice = input(f"\n收藏哪一个？输入数字(1-{len(matches)})，输入0取消收藏/直接输入要查询的术语：").strip().lower()
                if choice.isdigit():
                    idx = int(choice)
                    if idx == 0:
                        return None
                    if 1 <= idx <= len(matches):
                        term = matches[idx-1]
                        if term in self.favorites:
                            print(f"\n'{term}' 已在收藏夹中")
                            return None
                        self.favorites[term] = self.music_terms[term]
                        if self.save_favorites():
                            print("收藏成功！\n")
                        else:
                            print("收藏失败，请检查文件权限\n")
                        return None
                    else:
                        print(f"请输入1-{len(matches)}之间的数字！")
                else:
                    return choice  # 将非数字的输入作为新的搜索词

    def show_statistics(self):
        """显示词典统计信息"""
        print("\n词典统计：")
        print(f"- 总术语数：{len(self.music_terms)}")
        print(f"- 收藏数量：{len(self.favorites)}")
        print(f"- 搜索历史：{len(self.search_history)}")
        
        # 统计最常搜索的术语
        if self.search_history:
            most_common = Counter(self.search_history).most_common(5)
            print("\n最常搜索的术语：")
            for term, count in most_common:
                print(f"- {term}（{count}次）")
        print()

    def backup_data(self):
        """备份用户数据和程序"""
        try:
            import time
            import shutil
            import zipfile
            
            # 创建备份目录
            backup_dir = os.path.join(self.data_dir, 'backups')
            os.makedirs(backup_dir, exist_ok=True)
            
            # 生成备份文件名（使用时间戳）
            timestamp = time.strftime('%Y%m%d_%H%M%S')
            
            # 1. 备份用户数据
            data_backup_file = os.path.join(backup_dir, f'data_backup_{timestamp}.json')
            backup_data = {
                'favorites': self.favorites,
                'history': self.search_history,
                'timestamp': timestamp,
                'version': self.version
            }
            
            with open(data_backup_file, 'w', encoding='utf-8') as f:
                json.dump(backup_data, f, ensure_ascii=False, indent=4)
                
            # 2. 备份程序文件
            program_file = os.path.abspath(__file__)
            program_backup_file = os.path.join(backup_dir, f'program_backup_{timestamp}.zip')
            
            with zipfile.ZipFile(program_backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
                zipf.write(program_file, os.path.basename(program_file))
                
            print(f"\n数据已成功备份到：{data_backup_file}")
            print(f"程序已成功备份到：{program_backup_file}")
            return True
            
        except Exception as e:
            print(f"\n备份失败：{str(e)}")
            return False

    def restore_backup(self):
        """恢复备份数据（仅恢复用户数据，不恢复程序）"""
        backup_dir = os.path.join(self.data_dir, 'backups')
        if not os.path.exists(backup_dir):
            print("\n未找到备份文件")
            return False
            
        # 获取所有数据备份文件
        backups = sorted([f for f in os.listdir(backup_dir) if f.startswith('data_backup_')])
        if not backups:
            print("\n未找到备份文件")
            return False
            
        print("\n可用的备份：")
        for i, backup in enumerate(backups, 1):
            # 从文件名中提取时间信息
            timestamp = backup.replace('data_backup_', '').replace('.json', '')
            formatted_time = f"{timestamp[:4]}-{timestamp[4:6]}-{timestamp[6:8]} {timestamp[9:11]}:{timestamp[11:13]}:{timestamp[13:15]}"
            print(f"{i}. {formatted_time}")
            
        while True:
            choice = input("\n请选择要恢复的备份编号（0取消）：").strip()
            if choice == '0':
                print("\n已取消恢复")
                return False
                
            if choice.isdigit() and 1 <= int(choice) <= len(backups):
                try:
                    backup_file = os.path.join(backup_dir, backups[int(choice)-1])
                    with open(backup_file, 'r', encoding='utf-8') as f:
                        backup_data = json.load(f)
                    
                    # 确认恢复
                    confirm = input("\n恢复将覆盖当前数据，是否继续？(y/n): ").strip().lower()
                    if confirm != 'y':
                        print("\n已取消恢复")
                        return False
                    
                    # 恢复数据
                    self.favorites = backup_data['favorites']
                    self.search_history = backup_data['history']
                    self.save_favorites()
                    self.save_history()
                    
                    print("\n数据已成功恢复")
                    return True
                except Exception as e:
                    print(f"\n恢复失败：{str(e)}")
                    return False
            else:
                print("请输入有效的编号！")

    def _check_updates(self):
        """检查更新"""
        # 这里可以添加检查更新的逻辑
        pass

    def show_help(self):
        """显示帮助信息"""
        help_text = """
使用说明：
1. 搜索方式：
   - 支持中文搜索：如"甜美"
   - 支持拼音搜索：如"tianmei"
   - 支持英文搜索：如"dolce"
   - 支持部分匹配：如"dol"

2. 常用命令：
   - ? : 显示此帮助信息
   - dc1 : 退出程序
   - cl1 : 清空收藏夹
   - exp : 导出收藏夹
   - fav : 查看收藏夹
   - his : 查看搜索历史
   - stat : 显示统计信息

3. 搜索技巧：
   - 可以使用拼音首字母：如"tm"代替"tianmei"
   - 优先使用完整的关键词
   - 如果结果太多，尝试使用更长的关键词
   - 如果找不到，尝试使用同义词

4. 收藏功能：
   - 搜索到单个术语时，可以直接选择是否收藏
   - 搜索到多个术语时，输入编号选择要收藏的术语
   - 输入0可以取消收藏操作
"""
        return help_text

def check_duplicates():
    """检查音乐术语字典中的英语术语重复项"""
    # 创建字典来存储术语及其所有出现的行号
    term_locations = {}  # 只存储英文术语位置
    duplicates_found = False
    
    # 获取当前文件内容
    with open(__file__, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # 找到音乐术语字典的起始和结束位置
    start_line = 0
    end_line = 0
    for i, line in enumerate(lines, 1):
        if "music_terms = {" in line:
            start_line = i
        elif start_line > 0 and "}" in line and "music_terms.update" not in line:
            end_line = i
            break
    
    # 检查重复项
    for line_num in range(start_line, end_line + 1):
        line = lines[line_num - 1]
        if ':' in line and ("'" in line or '"' in line):
            try:
                # 只提取英文术语部分
                term = line.split(':')[0].strip().strip("'").strip('"').strip()
                
                # 标准化处理（转小写）
                normalized_term = term.lower()
                
                # 检查是否为英文术语（排除中文和空行）
                if normalized_term and all(ord(c) < 128 for c in normalized_term):
                    if normalized_term in term_locations:
                        term_locations[normalized_term]['lines'].append(line_num)
                        duplicates_found = True
                    else:
                        term_locations[normalized_term] = {
                            'value': term,
                            'lines': [line_num]
                        }
                        
            except Exception as e:
                continue
    
    if duplicates_found:
        print("\n发现英语术语重复：")
        
        # 显示重复项
        for term, info in term_locations.items():
            if len(info['lines']) > 1:
                print(f"'{info['value']}' 出现在以下行：{info['lines']}")
        
        response = input("\n是否要删除重复项？(y/n): ").strip().lower()
        if response == 'y':
            # 收集需要删除的行号
            lines_to_delete = set()
            
            # 对于每个重复术语，保留第一次出现的，删除其他所有出现
            for info in term_locations.values():
                if len(info['lines']) > 1:
                    # 保留第一次出现的，删除后面的所有重复
                    lines_to_delete.update(info['lines'][1:])
            
            # 转换为列表并排序（从大到小）
            lines_to_delete = sorted(list(lines_to_delete), reverse=True)
            
            # 删除重复行
            for line_num in lines_to_delete:
                del lines[line_num - 1]
            
            # 保存修改后的文件
            try:
                with open(__file__, 'w', encoding='utf-8') as f:
                    f.writelines(lines)
                print(f"已删除 {len(lines_to_delete)} 个重复的英语术语")
            except Exception as e:
                print(f"保存文件时发生错误：{str(e)}")
                return
    else:
        print("未发现英语术语重复，程序正常启动...\n")

def show_backup_tip():
    """显示备份提示信息"""
    tip = """
Tips: 如果程序文件丢失，可以通过以下步骤恢复：
1. 在程序原本所在的位置按 Command + Shift + . 显示隐藏目录
2. 进入 .music_dict_data/backups 目录
3. 找到最新的 program_backup_*.zip 文件
4. 解压该文件获取程序备份
"""
    print(tip)

def main():
    dictionary = MusicDictionary()
    # 启动时显示提示
    show_backup_tip()
    
    print("欢迎使用音乐术语词典！由yuze研发，完全开源。")
    print("使用说明：")
    print("- 输入 'dc1' 退出程序")
    print("- 输入 'cl1' 查看收藏")
    print("- 输入 'rm' 移除收藏")
    print("- 输入 'help' 显示帮助信息")
    print("- 输入 'exp' 导出收藏夹")
    print("- 输入 'his' 查看搜索历史（最多20条）")
    print("- 输入 'tst' 测试收藏夹术语")
    print("- 输入 'stat' 显示统计信息")
    print("- 输入 'backup' 备份数据")
    print("- 输入 'restore' 恢复备份")
    print("- 支持中文、英文、拼音搜索\n")
    
    while True:
        try:
            term = input("请输入音乐术语或中文解释：").strip()
            
            if not term:  # 处理空输入
                print("请输入要查询的内容\n")
                continue
            
            # 处理命令
            if term.lower() == 'dc1':
                print("谢谢使用，尊敬的音乐家，愿您在音乐的道路上继续前行。")
                dictionary.save_favorites()
                # 退出前显示提示
                show_backup_tip()
                break
            elif term.lower() == 'cl1':
                dictionary.show_favorites()
                continue
            elif term.lower() == 'rm':
                dictionary.remove_favorite()
                continue
            elif term.lower() == 'exp':
                dictionary.export_favorites()
                continue
            elif term.lower() == 'his':
                dictionary.show_history()
                continue
            elif term.lower() == 'tst':
                dictionary.favorites_quiz()
                continue
            elif term.lower() == 'stat':
                dictionary.show_statistics()
                continue
            elif term.lower() == 'backup':
                dictionary.backup_data()
                continue
            elif term.lower() == 'restore':
                dictionary.restore_backup()
                continue
            elif term.lower() == 'help':
                print("\n使用说明：")
                print("- 输入 'dc1' 退出程序")
                print("- 输入 'cl1' 查看收藏")
                print("- 输入 'rm' 移除收藏")
                print("- 输入 'help' 显示此帮助信息")
                print("- 输入 'exp' 导出收藏夹")
                print("- 输入 'his' 查看搜索历史")
                print("- 输入 'tst' 测试收藏夹术语")
                print("- 输入 'stat' 显示统计信息")
                print("- 输入 'backup' 备份数据")
                print("- 输入 'restore' 恢复备份")
                print("- 支持中文、英文、拼音搜索\n")
                continue
            
            # 查询术语
            result, original_term, matches = dictionary.query_term(term)
            print(f"\n{result}")
            
            # 记录搜索历史
            dictionary.add_to_history(term)
            
            # 处理匹配结果，获取可能的新搜索词
            new_term = dictionary.handle_search_result(original_term, matches)
            if new_term:  # 如果返回了新搜索词
                term = new_term
                result, original_term, matches = dictionary.query_term(term)
                print(f"\n{result}")
                dictionary.add_to_history(term)
            
            print()  # 添加空行分隔

        except Exception as e:
            print(f"发生错误：{str(e)}")
            print("请重试\n")

if __name__ == "__main__":
    check_duplicates()
    main()
